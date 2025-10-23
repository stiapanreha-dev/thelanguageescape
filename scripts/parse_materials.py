"""
Script to parse course materials from docs/Материалы/По дням/
Extracts tasks from DOCX files and creates JSON structure
"""
import json
import os
import sys
from pathlib import Path
from docx import Document

# Add parent directory to path
sys.path.append(str(Path(__file__).parent.parent))

MATERIALS_DIR = Path(__file__).parent.parent / "docs" / "Материалы" / "По дням"
OUTPUT_DIR = Path(__file__).parent.parent / "materials"


def parse_task_from_docx(docx_path):
    """Parse tasks from DOCX file"""
    try:
        doc = Document(docx_path)
        content = []

        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    content.append(row_text)

        return '\n'.join(content)

    except Exception as e:
        print(f"Error parsing {docx_path}: {e}")
        return None


def extract_tasks_from_text(text, day_number):
    """Extract structured tasks from parsed text"""
    tasks = []

    # This is a simplified parser - you'll need to customize based on actual format
    lines = text.split('\n')

    current_task = None
    task_number = 0

    for line in lines:
        line = line.strip()

        # Detect task headers
        if any(keyword in line.lower() for keyword in ['визуальный тест', 'голосовой', 'диалог', 'задание']):
            if current_task:
                tasks.append(current_task)

            task_number += 1
            task_type = 'choice'

            if 'голосов' in line.lower():
                task_type = 'voice'
            elif 'диалог' in line.lower():
                task_type = 'dialog'

            current_task = {
                'day': day_number,
                'task_number': task_number,
                'type': task_type,
                'title': line,
                'question': '',
                'options': [],
                'correct_answer': '',
                'hints': [],
                'voice_keywords': [],
                'dialog_steps': []
            }

        # Extract options (A/B/C/D format)
        elif current_task and line.startswith('🔘'):
            option = line.replace('🔘', '').strip()
            is_correct = '✅' in option
            option = option.replace('✅', '').strip()

            current_task['options'].append(option)
            if is_correct:
                current_task['correct_answer'] = option

        # Extract question
        elif current_task and 'бот:' in line.lower():
            current_task['question'] = line.split(':', 1)[1].strip().strip('"')

        # Add to description
        elif current_task and line:
            if not current_task['question']:
                current_task['question'] += ' ' + line

    # Add last task
    if current_task:
        tasks.append(current_task)

    return tasks


def scan_day_materials(day_path, day_number):
    """Scan materials for a specific day"""
    ready_path = day_path / "Готовое"

    if not ready_path.exists():
        print(f"Warning: {ready_path} does not exist")
        return None

    day_data = {
        'day': day_number,
        'title': f'Day {day_number}',
        'video': None,
        'brief': None,
        'audio': [],
        'images': [],
        'tasks': []
    }

    # Scan files
    for file in ready_path.iterdir():
        if file.is_file():
            filename = file.name.lower()

            # Video files
            if filename.endswith('.mp4'):
                if not day_data['video']:  # Take first video
                    day_data['video'] = str(file.relative_to(Path(__file__).parent.parent))

            # PDF briefs
            elif filename.endswith('.pdf'):
                day_data['brief'] = str(file.relative_to(Path(__file__).parent.parent))
                # Extract title from PDF name
                title = file.stem.replace('Protocol', '').replace('Day', '').strip()
                if title:
                    day_data['title'] = title

            # Audio files
            elif filename.endswith(('.mp3', '.wav', '.ogg')):
                day_data['audio'].append(str(file.relative_to(Path(__file__).parent.parent)))

            # Images
            elif filename.endswith(('.jpg', '.jpeg', '.png')):
                day_data['images'].append(str(file.relative_to(Path(__file__).parent.parent)))

            # Task DOCX files
            elif filename.startswith('задания') and filename.endswith('.docx'):
                task_text = parse_task_from_docx(file)
                if task_text:
                    tasks = extract_tasks_from_text(task_text, day_number)
                    day_data['tasks'] = tasks

    return day_data


def parse_all_materials():
    """Parse all course materials"""
    if not MATERIALS_DIR.exists():
        print(f"Error: Materials directory not found: {MATERIALS_DIR}")
        return None

    course_data = {
        'course_name': 'The Language Escape',
        'total_days': 10,
        'liberation_code': 'LIBERATION',
        'days': []
    }

    # Scan each day
    for day_num in range(1, 11):
        day_folder = MATERIALS_DIR / f"день{day_num:02d}"

        if not day_folder.exists():
            print(f"Warning: Day {day_num} folder not found: {day_folder}")
            continue

        print(f"Parsing day {day_num}...")
        day_data = scan_day_materials(day_folder, day_num)

        if day_data:
            course_data['days'].append(day_data)
            print(f"  ✅ Day {day_num}: {len(day_data['tasks'])} tasks found")
        else:
            print(f"  ❌ Day {day_num}: No data")

    return course_data


def save_course_data(course_data):
    """Save parsed data to JSON"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    output_file = OUTPUT_DIR / "course_data.json"

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(course_data, f, ensure_ascii=False, indent=2)

    print(f"\n✅ Course data saved to: {output_file}")
    print(f"Total days parsed: {len(course_data['days'])}")

    # Save each day separately
    for day_data in course_data['days']:
        day_file = OUTPUT_DIR / f"day_{day_data['day']:02d}.json"
        with open(day_file, 'w', encoding='utf-8') as f:
            json.dump(day_data, f, ensure_ascii=False, indent=2)

    print(f"Individual day files saved to: {OUTPUT_DIR}")


def main():
    """Main function"""
    print("=" * 60)
    print("Parsing course materials...")
    print("=" * 60)

    course_data = parse_all_materials()

    if course_data:
        save_course_data(course_data)
        print("\n✅ Parsing completed successfully!")
        print(f"\nSummary:")
        print(f"  - Total days: {len(course_data['days'])}")
        total_tasks = sum(len(day['tasks']) for day in course_data['days'])
        print(f"  - Total tasks: {total_tasks}")
    else:
        print("\n❌ Parsing failed!")
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())
